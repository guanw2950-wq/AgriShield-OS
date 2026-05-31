"""
NDVI crop-growth analysis service.

The API passes boundary/NDVI inputs and an output directory; this module returns
frontend-ready URLs, statistics, map artifacts, and a templated DOCX report.
"""

from __future__ import annotations

import csv
import json
import math
import os
import shutil
import zipfile
from pathlib import Path
from typing import Any


LEVEL_LABELS_5 = {
    1: "差",
    2: "一般",
    3: "中",
    4: "良",
    5: "优",
}

LEVEL_COLORS_5 = {
    1: "#d64545",
    2: "#ef8f35",
    3: "#e4c441",
    4: "#8cc152",
    5: "#2f9e44",
}


def _same_crs(left: Any, right: Any) -> bool:
    if left is None or right is None:
        return False
    try:
        from pyproj import CRS

        return CRS.from_user_input(left).equals(CRS.from_user_input(right), ignore_axis_order=True)
    except Exception:
        return str(left) == str(right)


def _infer_boundary_crs(bounds: Any) -> str | None:
    minx, miny, maxx, maxy = [float(value) for value in bounds]
    if not all(math.isfinite(value) for value in (minx, miny, maxx, maxy)):
        return None
    if -180 <= minx <= 180 and -180 <= maxx <= 180 and -90 <= miny <= 90 and -90 <= maxy <= 90:
        return "EPSG:4326"

    # CGCS2000 3-degree Gauss-Kruger zone coordinates often store the zone as
    # the leading two digits of the false easting, e.g. 38402868 -> zone 38.
    if 10_000_000 <= minx <= 60_000_000 and 0 <= miny <= 10_000_000:
        zone = int(minx // 1_000_000)
        if 25 <= zone <= 45:
            return f"EPSG:{4488 + zone}"

    return None


def _read_boundary_gdf(
    boundary_path: Path,
    work_dir: Path,
    target_crs: Any | None = None,
    source_crs: str | None = None,
):
    """Read GeoJSON/GPKG/KML/SHP zip and optionally project to target CRS."""
    import geopandas as gpd

    read_path = boundary_path
    if boundary_path.suffix.lower() == ".zip":
        extract_dir = work_dir / "boundary_unzipped"
        if extract_dir.exists():
            shutil.rmtree(extract_dir)
        extract_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(boundary_path) as zf:
            zf.extractall(extract_dir)
        shp_files = list(extract_dir.rglob("*.shp"))
        if not shp_files:
            raise ValueError("边界 ZIP 中没有找到 .shp 文件")
        read_path = shp_files[0]

    os.environ.setdefault("SHAPE_RESTORE_SHX", "YES")
    gdf = gpd.read_file(read_path)
    if gdf.empty:
        raise ValueError("边界文件为空")
    gdf = gdf[gdf.geometry.notna()].copy()
    if gdf.empty:
        raise ValueError("边界文件没有有效几何")
    try:
        gdf["geometry"] = gdf.geometry.make_valid()
    except Exception:
        gdf["geometry"] = gdf.buffer(0)
    if source_crs:
        gdf = gdf.set_crs(source_crs, allow_override=True)
    elif gdf.crs is None:
        inferred_crs = _infer_boundary_crs(gdf.total_bounds)
        if not inferred_crs:
            raise ValueError("边界文件缺少 CRS，请上传 .prj 或在请求中提供 boundary_crs")
        gdf = gdf.set_crs(inferred_crs, allow_override=True)

    if target_crs is not None and not _same_crs(gdf.crs, target_crs):
        gdf = gdf.to_crs(target_crs)
    return gdf


def calculate_boundary_area_mu(boundary_path: Path, work_dir: Path, source_crs: str | None = None) -> float:
    """Calculate boundary area in mu from the uploaded vector geometry."""
    gdf = _read_boundary_gdf(boundary_path, work_dir, source_crs=source_crs)
    if gdf.crs is None:
        raise ValueError("边界文件缺少 CRS，无法计算面积")

    try:
        is_geographic = bool(gdf.crs.is_geographic)
    except Exception:
        is_geographic = False

    if is_geographic:
        area_crs = gdf.estimate_utm_crs() or "EPSG:6933"
        area_gdf = gdf.to_crs(area_crs)
    else:
        area_gdf = gdf

    area_sqm = float(area_gdf.geometry.area.sum())
    if not math.isfinite(area_sqm) or area_sqm <= 0:
        raise ValueError("边界面积计算结果无效")
    return round(area_sqm * 0.0015, 2)


def create_synthetic_ndvi_raster(
    boundary_path: Path,
    work_dir: Path,
    out_tif: Path,
    source_crs: str | None = None,
) -> dict[str, Any]:
    """Create a deterministic NDVI raster for local/dev runs when GEE is unavailable."""
    import numpy as np
    import rasterio
    from rasterio.transform import from_origin

    gdf = _read_boundary_gdf(boundary_path, work_dir, source_crs=source_crs)
    if gdf.crs is None:
        raise ValueError("边界文件缺少 CRS，无法生成 NDVI 栅格")

    try:
        is_geographic = bool(gdf.crs.is_geographic)
    except Exception:
        is_geographic = False

    raster_gdf = gdf.to_crs(gdf.estimate_utm_crs() or "EPSG:3857") if is_geographic else gdf
    minx, miny, maxx, maxy = [float(value) for value in raster_gdf.total_bounds]
    span = max(maxx - minx, maxy - miny)
    if not math.isfinite(span) or span <= 0:
        raise ValueError("边界范围无效，无法生成 NDVI 栅格")

    # Keep very small sample plots readable while capping larger scenes.
    resolution = max(2.0, span / 480.0)
    pad = max(resolution * 8, span * 0.08)
    width = max(16, int(math.ceil((maxx - minx + pad * 2) / resolution)))
    height = max(16, int(math.ceil((maxy - miny + pad * 2) / resolution)))
    if max(width, height) > 900:
        scale = max(width, height) / 900.0
        width = max(16, int(width / scale))
        height = max(16, int(height / scale))
        resolution *= scale

    x = np.linspace(0, 1, width, dtype="float32")
    y = np.linspace(0, 1, height, dtype="float32")[:, None]
    seed = int(abs(minx * 11 + miny * 17 + maxx * 23 + maxy * 31)) % (2**32 - 1)
    rng = np.random.default_rng(seed)
    ndvi = (
        0.24
        + 0.46 * x
        + 0.12 * np.sin(y * np.pi * 2.4)
        + 0.07 * np.cos((x + y) * np.pi * 3.0)
        + rng.normal(0, 0.03, (height, width)).astype("float32")
    )
    ndvi = np.clip(ndvi, 0.03, 0.88).astype("float32")

    out_tif.parent.mkdir(parents=True, exist_ok=True)
    transform = from_origin(minx - pad, maxy + pad, resolution, resolution)
    with rasterio.open(
        out_tif,
        "w",
        driver="GTiff",
        height=height,
        width=width,
        count=1,
        dtype="float32",
        crs=raster_gdf.crs,
        transform=transform,
        nodata=-9999.0,
        compress="lzw",
    ) as dst:
        dst.write(ndvi, 1)

    return {
        "source": "synthetic",
        "source_label": "本地模拟 NDVI",
        "width": width,
        "height": height,
        "crs": str(raster_gdf.crs),
        "resolution": round(float(resolution), 4),
        "note": "未使用真实卫星影像，仅用于本地自动流程验证",
    }


def export_gee_ndvi_raster(
    boundary_path: Path,
    work_dir: Path,
    out_tif: Path,
    source_crs: str | None = None,
    start_date: str = "2025-08-30",
    end_date: str = "2025-09-15",
    max_cloud_pct: float = 30.0,
    scale: int = 10,
    project_id: str | None = None,
) -> dict[str, Any]:
    """Export Sentinel-2 NDVI from Google Earth Engine into a local GeoTIFF."""
    import requests
    from shapely.geometry import mapping
    from shapely.ops import unary_union

    try:
        import ee
    except Exception as exc:
        raise RuntimeError("未安装 earthengine-api，无法自动获取 GEE Sentinel-2 NDVI") from exc

    try:
        if project_id:
            ee.Initialize(project=project_id)
        else:
            ee.Initialize()
    except Exception as exc:
        raise RuntimeError("GEE 未初始化，请先完成 Earth Engine 认证或改用演示数据源") from exc

    gdf = _read_boundary_gdf(boundary_path, work_dir, source_crs=source_crs)
    if not _same_crs(gdf.crs, "EPSG:4326"):
        gdf = gdf.to_crs("EPSG:4326")
    geometry = mapping(unary_union(gdf.geometry))
    roi = ee.Geometry(geometry)

    def mask_s2_sr(image):
        scl = image.select("SCL")
        clear = (
            scl.neq(3)
            .And(scl.neq(8))
            .And(scl.neq(9))
            .And(scl.neq(10))
            .And(scl.neq(11))
        )
        return image.updateMask(clear)

    def add_ndvi(image):
        return image.addBands(image.normalizedDifference(["B8", "B4"]).rename("ndvi"))

    collection = (
        ee.ImageCollection("COPERNICUS/S2_SR_HARMONIZED")
        .filterDate(start_date, end_date)
        .filterBounds(roi)
        .filter(ee.Filter.lt("CLOUDY_PIXEL_PERCENTAGE", max_cloud_pct))
    )
    image_count = int(collection.size().getInfo())
    if image_count == 0:
        raise RuntimeError(f"{start_date} 至 {end_date} 内没有满足云量阈值的 Sentinel-2 影像")

    ndvi = collection.map(mask_s2_sr).map(add_ndvi).median().select("ndvi").clip(roi)
    download_url = ndvi.getDownloadURL(
        {
            "scale": scale,
            "region": geometry,
            "filePerBand": False,
            "format": "GEO_TIFF",
        }
    )

    response = requests.get(download_url, timeout=180)
    response.raise_for_status()
    out_tif.parent.mkdir(parents=True, exist_ok=True)
    content_type = response.headers.get("content-type", "")
    if "zip" in content_type or response.content[:2] == b"PK":
        archive_path = work_dir / "gee_ndvi_download.zip"
        archive_path.write_bytes(response.content)
        extract_dir = work_dir / "gee_ndvi_download"
        if extract_dir.exists():
            shutil.rmtree(extract_dir)
        extract_dir.mkdir(parents=True, exist_ok=True)
        with zipfile.ZipFile(archive_path) as zf:
            zf.extractall(extract_dir)
        tif_files = list(extract_dir.rglob("*.tif")) + list(extract_dir.rglob("*.tiff"))
        if not tif_files:
            raise RuntimeError("GEE 下载结果中没有 GeoTIFF")
        shutil.copy2(tif_files[0], out_tif)
    else:
        out_tif.write_bytes(response.content)

    return {
        "source": "gee",
        "source_label": "GEE Sentinel-2 SR NDVI",
        "start_date": start_date,
        "end_date": end_date,
        "max_cloud_pct": max_cloud_pct,
        "scale": scale,
        "image_count": image_count,
    }


def prepare_auto_ndvi_raster(
    boundary_path: Path,
    work_dir: Path,
    out_tif: Path,
    source_crs: str | None = None,
    ndvi_source: str = "auto",
    start_date: str = "2025-08-30",
    end_date: str = "2025-09-15",
    max_cloud_pct: float = 30.0,
    gee_project_id: str | None = None,
) -> dict[str, Any]:
    ndvi_source = (ndvi_source or "auto").lower()
    if ndvi_source not in {"auto", "gee", "synthetic"}:
        raise ValueError("ndvi_source 只能为 auto、gee 或 synthetic")

    if ndvi_source in {"auto", "gee"}:
        try:
            return export_gee_ndvi_raster(
                boundary_path=boundary_path,
                work_dir=work_dir,
                out_tif=out_tif,
                source_crs=source_crs,
                start_date=start_date,
                end_date=end_date,
                max_cloud_pct=max_cloud_pct,
                project_id=gee_project_id,
            )
        except Exception as exc:
            if ndvi_source == "gee":
                raise
            synthetic = create_synthetic_ndvi_raster(boundary_path, work_dir, out_tif, source_crs)
            synthetic["fallback_reason"] = str(exc)
            return synthetic

    return create_synthetic_ndvi_raster(boundary_path, work_dir, out_tif, source_crs)


def clip_raster_by_boundary(
    src_tif: Path,
    boundary_path: Path,
    out_tif: Path,
    work_dir: Path,
    source_crs: str | None = None,
) -> dict[str, Any]:
    """Clip the uploaded NDVI GeoTIFF by the uploaded boundary."""
    import rasterio
    from rasterio.mask import mask
    from shapely.geometry import mapping

    with rasterio.open(src_tif) as src:
        boundary_gdf = _read_boundary_gdf(boundary_path, work_dir, src.crs, source_crs)
        geometries = [mapping(geom) for geom in boundary_gdf.geometry if not geom.is_empty]
        if not geometries:
            raise ValueError("边界文件没有可用于裁剪的几何")

        nodata = src.nodata
        if nodata is None:
            nodata = 0
        clipped, transform = mask(src, geometries, crop=True, nodata=nodata, filled=True)
        meta = src.meta.copy()
        meta.update(
            {
                "height": clipped.shape[1],
                "width": clipped.shape[2],
                "transform": transform,
                "nodata": nodata,
                "compress": "lzw",
            }
        )

    out_tif.parent.mkdir(parents=True, exist_ok=True)
    with rasterio.open(out_tif, "w", **meta) as dst:
        dst.write(clipped)

    return {
        "width": int(meta["width"]),
        "height": int(meta["height"]),
        "crs": str(meta.get("crs") or ""),
        "nodata": nodata,
    }


def iter_raster_windows(width: int, height: int, tile_size: int = 1024):
    from rasterio.windows import Window

    for top in range(0, height, tile_size):
        rows = min(tile_size, height - top)
        for left in range(0, width, tile_size):
            cols = min(tile_size, width - left)
            yield Window(left, top, cols, rows)


def sample_raster_values(
    raster_path: Path,
    tile_size: int = 1024,
    max_samples: int = 2_000_000,
    per_chunk_limit: int = 250_000,
    seed: int = 2024,
) -> tuple[Any, dict[str, Any]]:
    """Collect representative valid NDVI samples without loading the whole raster."""
    import numpy as np
    import rasterio

    rng = np.random.default_rng(seed)
    samples = []
    total_sampled = 0
    stats: dict[str, Any] = {"min": None, "max": None, "count": 0}

    with rasterio.open(raster_path) as src:
        nodata = src.nodata
        for window in iter_raster_windows(src.width, src.height, tile_size):
            data = src.read(1, window=window, masked=False)
            mask = np.isfinite(data)
            if nodata is not None:
                mask &= data != nodata
            if not mask.any():
                continue

            valid = data[mask].astype("float32", copy=False)
            stats["count"] += int(valid.size)
            chunk_min = float(valid.min())
            chunk_max = float(valid.max())
            stats["min"] = chunk_min if stats["min"] is None else min(stats["min"], chunk_min)
            stats["max"] = chunk_max if stats["max"] is None else max(stats["max"], chunk_max)

            if per_chunk_limit and valid.size > per_chunk_limit:
                valid = valid[rng.choice(valid.size, size=per_chunk_limit, replace=False)]

            samples.append(valid)
            total_sampled += valid.size

            if total_sampled > max_samples * 2:
                merged = np.concatenate(samples)
                merged = merged[rng.choice(merged.size, size=max_samples, replace=False)]
                samples = [merged]
                total_sampled = merged.size

    if not samples:
        raise ValueError("NDVI 有效像元为空")

    merged = np.concatenate(samples)
    if merged.size > max_samples:
        merged = merged[rng.choice(merged.size, size=max_samples, replace=False)]
    stats["sample_size"] = int(merged.size)
    return merged, stats


def compute_class_breaks(values: Any, method: str, n_classes: int, stats: dict[str, Any]) -> list[float]:
    import numpy as np

    if n_classes < 2 or n_classes > 9:
        raise ValueError("n_classes 需在 2 到 9 之间")

    vals = values[np.isfinite(values)]
    if vals.size == 0:
        raise ValueError("无法从空样本计算分级阈值")

    method = method.lower()
    min_val = float(stats["min"] if stats.get("min") is not None else np.nanmin(vals))
    max_val = float(stats["max"] if stats.get("max") is not None else np.nanmax(vals))
    if math.isclose(min_val, max_val):
        return [max_val for _ in range(n_classes)]

    try:
        import mapclassify as mc

        if method == "jenks":
            return [float(x) for x in mc.NaturalBreaks(vals, k=n_classes).bins.tolist()]
        if method == "equalinterval":
            return [float(x) for x in mc.EqualInterval(vals, k=n_classes).bins.tolist()]
        if method == "quantile":
            return [float(x) for x in mc.Quantiles(vals, k=n_classes).bins.tolist()]
        if method == "std":
            return [float(x) for x in mc.StdMean(vals, multiples=[-2, -1, 0, 1, 2]).bins.tolist()[:n_classes]]
    except Exception:
        # Fall through to a deterministic numpy implementation.
        pass

    if method == "equalinterval":
        return [float(x) for x in np.linspace(min_val, max_val, n_classes + 1)[1:]]

    # Quantile fallback is also used when Jenks/mapclassify is unavailable.
    quantiles = np.linspace(0, 1, n_classes + 1)[1:]
    bins = np.quantile(vals, quantiles).astype(float).tolist()
    bins[-1] = max_val
    return [float(x) for x in bins]


def classify_block(data: Any, bins_full: list[float], nodata: float | int | None):
    import numpy as np

    arr = data.astype("float32", copy=False)
    mask = np.isfinite(arr)
    if nodata is not None:
        mask &= arr != nodata
    result = np.zeros(arr.shape, dtype="uint8")
    if mask.any():
        result[mask] = np.digitize(arr[mask], bins=bins_full, right=False).astype("uint8")
    return result


def write_classified_raster(
    in_ndvi: Path,
    out_tif: Path,
    class_breaks: list[float],
    tile_size: int = 1024,
) -> dict[int, int]:
    import numpy as np
    import rasterio

    bins_full = [-np.inf] + class_breaks[:-1] + [np.inf]
    count_map: dict[int, int] = {}

    with rasterio.open(in_ndvi) as src:
        meta = src.meta.copy()
        meta.update(count=1, dtype="uint8", nodata=0, compress="lzw")
        out_tif.parent.mkdir(parents=True, exist_ok=True)
        with rasterio.open(out_tif, "w", **meta) as dst:
            for window in iter_raster_windows(src.width, src.height, tile_size):
                data = src.read(1, window=window, masked=False)
                cls_block = classify_block(data, bins_full, src.nodata)
                dst.write(cls_block, 1, window=window)
                unique, counts = np.unique(cls_block, return_counts=True)
                for value, count in zip(unique.tolist(), counts.tolist()):
                    count_map[int(value)] = count_map.get(int(value), 0) + int(count)

    return count_map


def build_preview(dataset: Any, max_size: int = 2048):
    from rasterio.enums import Resampling

    height, width = dataset.height, dataset.width
    max_dim = max(height, width)
    if max_dim > max_size:
        scale = max_dim / float(max_size)
        out_h = max(1, int(height / scale))
        out_w = max(1, int(width / scale))
        return dataset.read(1, out_shape=(1, out_h, out_w), resampling=Resampling.nearest)[0]
    return dataset.read(1)


def render_ndvi_preview(in_ndvi: Path, out_png: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np
    import rasterio

    with rasterio.open(in_ndvi) as src:
        arr = build_preview(src)
        nodata = src.nodata

    arr = arr.astype("float32", copy=False)
    if nodata is not None:
        arr = np.ma.masked_equal(arr, nodata)

    try:
        lo = float(np.nanpercentile(arr, 2))
        hi = float(np.nanpercentile(arr, 98))
        if lo >= -1.5 and hi <= 1.5:
            lo, hi = -1.0, 1.0
    except Exception:
        lo, hi = None, None

    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(9, 7))
    im = ax.imshow(arr, cmap="RdYlGn", vmin=lo, vmax=hi)
    ax.set_axis_off()
    fig.colorbar(im, ax=ax, fraction=0.035, pad=0.02, label="NDVI")
    fig.tight_layout()
    fig.savefig(out_png, dpi=180, bbox_inches="tight")
    plt.close(fig)


def render_class_preview(class_tif: Path, out_png: Path, n_classes: int = 5) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.patches as mpatches
    import matplotlib.pyplot as plt
    import numpy as np
    import rasterio
    from matplotlib.colors import BoundaryNorm, ListedColormap

    if n_classes == 5:
        colors = [(1, 1, 1, 0)] + [_hex_to_rgba(LEVEL_COLORS_5[i]) for i in range(1, 6)]
        labels = [LEVEL_LABELS_5[i] for i in range(1, 6)]
    else:
        palette = plt.get_cmap("RdYlGn")
        colors = [(1, 1, 1, 0)] + [palette(i / max(n_classes - 1, 1)) for i in range(n_classes)]
        labels = [f"等级 {i}" for i in range(1, n_classes + 1)]

    with rasterio.open(class_tif) as src:
        data = build_preview(src)

    masked = np.where(data == 0, np.nan, data)
    cmap = ListedColormap(colors)
    norm = BoundaryNorm(np.arange(0, n_classes + 2) - 0.5, cmap.N)

    out_png.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(9, 7))
    ax.imshow(masked, cmap=cmap, norm=norm, aspect="equal")
    ax.set_axis_off()
    legend = [
        mpatches.Patch(color=colors[i], label=labels[i - 1])
        for i in range(1, n_classes + 1)
    ]
    ax.legend(handles=legend, loc="center left", bbox_to_anchor=(1.01, 0.5), frameon=False)
    fig.tight_layout()
    fig.savefig(out_png, dpi=180, bbox_inches="tight")
    plt.close(fig)


def raster_wgs84_bounds(raster_path: Path) -> list[float]:
    import rasterio
    from rasterio.warp import transform_bounds

    with rasterio.open(raster_path) as src:
        if src.crs is not None and not _same_crs(src.crs, "EPSG:4326"):
            bounds = transform_bounds(src.crs, "EPSG:4326", *src.bounds, densify_pts=21)
        else:
            bounds = src.bounds
    return [float(bounds[0]), float(bounds[1]), float(bounds[2]), float(bounds[3])]


def write_class_overlay_png(class_tif: Path, out_png: Path, n_classes: int = 5) -> list[float]:
    import numpy as np
    import rasterio
    from PIL import Image

    with rasterio.open(class_tif) as src:
        data = src.read(1, masked=False)

    rgba = np.zeros((data.shape[0], data.shape[1], 4), dtype=np.uint8)
    if n_classes == 5:
        color_map = LEVEL_COLORS_5
    else:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        palette = plt.get_cmap("RdYlGn")
        color_map = {
            idx: "#{:02x}{:02x}{:02x}".format(
                int(palette((idx - 1) / max(n_classes - 1, 1))[0] * 255),
                int(palette((idx - 1) / max(n_classes - 1, 1))[1] * 255),
                int(palette((idx - 1) / max(n_classes - 1, 1))[2] * 255),
            )
            for idx in range(1, n_classes + 1)
        }

    for value, color in color_map.items():
        rgb = [int(color.lstrip("#")[i : i + 2], 16) for i in (0, 2, 4)]
        mask = data == value
        rgba[mask, :3] = rgb
        rgba[mask, 3] = 218

    out_png.parent.mkdir(parents=True, exist_ok=True)
    Image.fromarray(rgba, mode="RGBA").save(out_png)
    return raster_wgs84_bounds(class_tif)


def write_ndvi_overlay_png(ndvi_tif: Path, out_png: Path) -> list[float]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.cm as cm
    import numpy as np
    import rasterio
    from PIL import Image

    with rasterio.open(ndvi_tif) as src:
        data = src.read(1, masked=False).astype("float32", copy=False)
        nodata = src.nodata

    valid = np.isfinite(data)
    if nodata is not None:
        valid &= data != nodata

    normalized = np.zeros(data.shape, dtype="float32")
    if valid.any():
        lo = float(np.nanpercentile(data[valid], 2))
        hi = float(np.nanpercentile(data[valid], 98))
        if lo >= -1.5 and hi <= 1.5:
            lo, hi = -0.2, 0.9
        if math.isclose(lo, hi):
            hi = lo + 1.0
        normalized[valid] = np.clip((data[valid] - lo) / (hi - lo), 0, 1)

    rgba = (cm.get_cmap("RdYlGn")(normalized) * 255).astype(np.uint8)
    rgba[~valid, 3] = 0
    rgba[valid, 3] = 224

    out_png.parent.mkdir(parents=True, exist_ok=True)
    Image.fromarray(rgba, mode="RGBA").save(out_png)
    return raster_wgs84_bounds(ndvi_tif)


def _hex_to_rgba(value: str) -> tuple[float, float, float, float]:
    value = value.lstrip("#")
    return (
        int(value[0:2], 16) / 255.0,
        int(value[2:4], 16) / 255.0,
        int(value[4:6], 16) / 255.0,
        1.0,
    )


def summarize_classes(
    count_map: dict[int, int],
    total_area_mu: float,
    n_classes: int,
) -> list[dict[str, Any]]:
    valid_total = sum(count for value, count in count_map.items() if value > 0)
    rows: list[dict[str, Any]] = []
    for value in range(1, n_classes + 1):
        count = int(count_map.get(value, 0))
        ratio = count / valid_total if valid_total else 0.0
        label = LEVEL_LABELS_5.get(value, f"等级 {value}") if n_classes == 5 else f"等级 {value}"
        color = LEVEL_COLORS_5.get(value, "#6b7280") if n_classes == 5 else "#6b7280"
        rows.append(
            {
                "value": value,
                "label": label,
                "count": count,
                "ratio": round(ratio, 6),
                "area_mu": round(ratio * float(total_area_mu), 2),
                "color": color,
            }
        )
    return rows


def write_summary_csv(rows: list[dict[str, Any]], out_csv: Path) -> None:
    out_csv.parent.mkdir(parents=True, exist_ok=True)
    with out_csv.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=["value", "label", "count", "ratio", "area_mu", "color"])
        writer.writeheader()
        writer.writerows(rows)


def write_class_geojson(class_tif: Path, out_geojson: Path, max_pixels: int = 8_000_000) -> Path | None:
    """Vectorize class raster for optional web-map overlays; skip very large rasters."""
    import geopandas as gpd
    import numpy as np
    import rasterio
    from rasterio.features import shapes
    from shapely.geometry import shape
    from shapely.ops import unary_union

    with rasterio.open(class_tif) as src:
        if src.width * src.height > max_pixels:
            return None
        data = src.read(1, masked=False)
        mask = data > 0
        if not mask.any():
            return None
        grouped: dict[int, list[Any]] = {}
        for geom, value in shapes(data, mask=mask, transform=src.transform):
            cls_id = int(value)
            grouped.setdefault(cls_id, []).append(shape(geom))

        features = [
            {"value": cls_id, "label": LEVEL_LABELS_5.get(cls_id, f"等级 {cls_id}"), "geometry": unary_union(geoms)}
            for cls_id, geoms in grouped.items()
            if geoms
        ]
        gdf = gpd.GeoDataFrame(features, geometry="geometry", crs=src.crs)

    if gdf.empty:
        return None
    if gdf.crs is not None and gdf.crs.to_epsg() != 4326:
        gdf = gdf.to_crs("EPSG:4326")
    out_geojson.parent.mkdir(parents=True, exist_ok=True)
    gdf.to_file(out_geojson, driver="GeoJSON")
    return out_geojson


def write_boundary_geojson(
    boundary_path: Path,
    work_dir: Path,
    out_geojson: Path,
    source_crs: str | None = None,
) -> tuple[Path, list[float]]:
    """Persist the uploaded boundary as WGS84 GeoJSON for the web map."""
    gdf = _read_boundary_gdf(boundary_path, work_dir, source_crs=source_crs)
    if gdf.crs is not None and not _same_crs(gdf.crs, "EPSG:4326"):
        gdf = gdf.to_crs("EPSG:4326")
    out_geojson.parent.mkdir(parents=True, exist_ok=True)
    gdf.to_file(out_geojson, driver="GeoJSON")
    minx, miny, maxx, maxy = [float(x) for x in gdf.total_bounds]
    return out_geojson, [minx, miny, maxx, maxy]


def _lonlat_to_global_pixel(lon: float, lat: float, zoom: int) -> tuple[float, float]:
    lat = max(min(lat, 85.05112878), -85.05112878)
    sin_lat = math.sin(math.radians(lat))
    size = 256 * (2**zoom)
    x = (lon + 180.0) / 360.0 * size
    y = (0.5 - math.log((1 + sin_lat) / (1 - sin_lat)) / (4 * math.pi)) * size
    return x, y


def _pick_report_zoom(bounds: list[float], width: int, height: int) -> int:
    minx, miny, maxx, maxy = bounds
    for zoom in range(19, 5, -1):
        x1, y1 = _lonlat_to_global_pixel(minx, maxy, zoom)
        x2, y2 = _lonlat_to_global_pixel(maxx, miny, zoom)
        if abs(x2 - x1) <= width * 0.70 and abs(y2 - y1) <= height * 0.70:
            return zoom
    return 6


def _fetch_satellite_canvas(
    bounds: list[float],
    out_size: tuple[int, int] = (690, 507),
) -> tuple[Any, dict[str, Any]]:
    from io import BytesIO

    import requests
    from PIL import Image, ImageDraw

    width, height = out_size
    view_zoom = _pick_report_zoom(bounds, width, height)
    native_zoom = min(view_zoom, 17)
    scale = 2 ** (view_zoom - native_zoom)
    center_lon = (bounds[0] + bounds[2]) / 2
    center_lat = (bounds[1] + bounds[3]) / 2
    center_x, center_y = _lonlat_to_global_pixel(center_lon, center_lat, view_zoom)
    top_left_view = (center_x - width / 2, center_y - height / 2)
    top_left_native = (top_left_view[0] / scale, top_left_view[1] / scale)
    native_width = math.ceil(width / scale)
    native_height = math.ceil(height / scale)

    min_tile_x = math.floor(top_left_native[0] / 256)
    min_tile_y = math.floor(top_left_native[1] / 256)
    max_tile_x = math.floor((top_left_native[0] + native_width) / 256)
    max_tile_y = math.floor((top_left_native[1] + native_height) / 256)
    tile_count = 2**native_zoom
    mosaic = Image.new(
        "RGB",
        ((max_tile_x - min_tile_x + 1) * 256, (max_tile_y - min_tile_y + 1) * 256),
        (42, 48, 39),
    )
    session = requests.Session()
    for tx in range(min_tile_x, max_tile_x + 1):
        for ty in range(min_tile_y, max_tile_y + 1):
            if ty < 0 or ty >= tile_count:
                continue
            wrapped_tx = tx % tile_count
            url = (
                "https://server.arcgisonline.com/ArcGIS/rest/services/"
                f"World_Imagery/MapServer/tile/{native_zoom}/{ty}/{wrapped_tx}"
            )
            try:
                response = session.get(url, timeout=8)
                response.raise_for_status()
                tile = Image.open(BytesIO(response.content)).convert("RGB")
            except Exception:
                tile = Image.new("RGB", (256, 256), (42, 48, 39))
                draw = ImageDraw.Draw(tile)
                draw.line([(0, 0), (255, 255)], fill=(68, 75, 63), width=1)
                draw.line([(255, 0), (0, 255)], fill=(68, 75, 63), width=1)
            mosaic.paste(tile, ((tx - min_tile_x) * 256, (ty - min_tile_y) * 256))

    crop_left = int(round(top_left_native[0] - min_tile_x * 256))
    crop_top = int(round(top_left_native[1] - min_tile_y * 256))
    crop = mosaic.crop((crop_left, crop_top, crop_left + native_width, crop_top + native_height))
    if crop.size != out_size:
        crop = crop.resize(out_size, Image.Resampling.BICUBIC)
    return crop.convert("RGBA"), {"zoom": view_zoom, "top_left": top_left_view, "size": out_size}


def _report_local_xy(lon: float, lat: float, view: dict[str, Any]) -> tuple[float, float]:
    x, y = _lonlat_to_global_pixel(lon, lat, int(view["zoom"]))
    top_left_x, top_left_y = view["top_left"]
    return x - top_left_x, y - top_left_y


def _draw_geojson_polygons(
    image: Any,
    geojson_data: dict[str, Any],
    view: dict[str, Any],
    *,
    fill: tuple[int, int, int, int],
    outline: tuple[int, int, int, int],
    width: int = 2,
) -> None:
    from PIL import Image
    from PIL import ImageDraw

    overlay = Image.new("RGBA", image.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay, "RGBA")

    def rings_from_geometry(geometry: dict[str, Any]):
        if geometry.get("type") == "Polygon":
            return geometry.get("coordinates", [])
        if geometry.get("type") == "MultiPolygon":
            for polygon in geometry.get("coordinates", []):
                for ring in polygon:
                    yield ring
            return
        return []

    for feature in geojson_data.get("features", []):
        geometry = feature.get("geometry") or {}
        if geometry.get("type") == "Polygon":
            rings = geometry.get("coordinates", [])
        elif geometry.get("type") == "MultiPolygon":
            rings = [ring for polygon in geometry.get("coordinates", []) for ring in polygon]
        else:
            rings = []
        if not rings:
            continue
        outer = [_report_local_xy(float(lon), float(lat), view) for lon, lat, *_ in rings[0]]
        if len(outer) >= 3:
            draw.polygon(outer, fill=fill)
            draw.line(outer + [outer[0]], fill=outline, width=width)

    image.alpha_composite(overlay)


def _overlay_report_raster(
    image: Any,
    overlay_png: Path,
    overlay_bounds: list[float],
    view: dict[str, Any],
    *,
    opacity: float,
    nearest: bool = True,
) -> None:
    from PIL import Image

    if not overlay_png.exists():
        return
    minx, miny, maxx, maxy = overlay_bounds
    left, top = _report_local_xy(minx, maxy, view)
    right, bottom = _report_local_xy(maxx, miny, view)
    box = (
        int(round(min(left, right))),
        int(round(min(top, bottom))),
        int(round(abs(right - left))),
        int(round(abs(bottom - top))),
    )
    if box[2] <= 1 or box[3] <= 1:
        return
    overlay = Image.open(overlay_png).convert("RGBA")
    resample = Image.Resampling.NEAREST if nearest else Image.Resampling.BILINEAR
    overlay = overlay.resize((box[2], box[3]), resample)
    if opacity < 1:
        alpha = overlay.getchannel("A").point(lambda value: int(value * opacity))
        overlay.putalpha(alpha)
    image.alpha_composite(overlay, dest=(box[0], box[1]))


def _report_font(size: int, bold: bool = False):
    from PIL import ImageFont

    candidates = [
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_report_legend(
    image: Any,
    title: str,
    rows: list[tuple[str, str, str | None]],
    *,
    left: int = 10,
    bottom: int = 10,
) -> None:
    from PIL import ImageDraw

    row_h = 22
    title_h = 24
    width = 112
    height = title_h + row_h * len(rows) + 10
    x0 = left
    y0 = image.height - bottom - height
    draw = ImageDraw.Draw(image, "RGBA")
    draw.rounded_rectangle((x0, y0, x0 + width, y0 + height), radius=4, fill=(223, 251, 255, 226))
    title_font = _report_font(16, True)
    row_font = _report_font(15)
    draw.text((x0 + 8, y0 + 6), title, fill=(0, 38, 44, 255), font=title_font)
    y = y0 + title_h + 5
    for label, color, extra in rows:
        rgb = tuple(int(color.lstrip("#")[i : i + 2], 16) for i in (0, 2, 4))
        draw.rectangle((x0 + 8, y + 4, x0 + 38, y + 18), fill=rgb + (255,), outline=(30, 30, 30, 160))
        text = label if extra is None else f"{label} {extra}"
        draw.text((x0 + 45, y + 1), text, fill=(0, 38, 44, 255), font=row_font)
        y += row_h


def create_report_map_images(
    *,
    boundary_geojson: Path,
    class_overlay_png: Path,
    ndvi_overlay_png: Path,
    overlay_bounds: list[float],
    out_growth_png: Path,
    out_crop_png: Path,
    summary: list[dict[str, Any]],
    crop_label: str,
) -> tuple[Path, Path]:
    """Render static report images: satellite + NDVI classification, satellite + SHP crop layer."""
    boundary_data = json.loads(boundary_geojson.read_text(encoding="utf-8"))
    bounds = overlay_bounds
    growth_img, view = _fetch_satellite_canvas(bounds, (690, 507))
    crop_img = growth_img.copy()

    _overlay_report_raster(growth_img, class_overlay_png, overlay_bounds, view, opacity=0.82, nearest=True)
    _draw_geojson_polygons(
        growth_img,
        boundary_data,
        view,
        fill=(0, 245, 255, 10),
        outline=(0, 255, 255, 230),
        width=2,
    )
    legend_rows = [
        (str(row["label"]), str(row["color"]), None)
        for row in sorted(summary, key=lambda item: item["value"], reverse=True)
    ]
    _draw_report_legend(growth_img, "长势等级", legend_rows, left=10, bottom=12)

    _draw_geojson_polygons(
        crop_img,
        boundary_data,
        view,
        fill=(32, 216, 210, 205),
        outline=(0, 230, 230, 245),
        width=2,
    )
    _draw_report_legend(crop_img, "图例", [(crop_label or "作物", "#20d8d2", None)], left=10, bottom=12)

    out_growth_png.parent.mkdir(parents=True, exist_ok=True)
    growth_img.convert("RGB").save(out_growth_png, quality=95)
    crop_img.convert("RGB").save(out_crop_png, quality=95)
    return out_growth_png, out_crop_png


def create_interactive_map_html(
    *,
    out_html: Path,
    boundary_geojson: Path,
    class_overlay_png: Path | None,
    ndvi_overlay_png: Path | None,
    summary: list[dict[str, Any]],
    bounds: list[float],
    raster_bounds: list[float] | None = None,
    crop_label: str = "作物",
) -> Path:
    """Create a Leaflet map using only locally generated analysis layers."""
    boundary_data = json.loads(boundary_geojson.read_text(encoding="utf-8"))
    legend_rows = "\n".join(
        f'<div class="legend-row"><i style="background:{row["color"]}"></i><span>{row["label"]}</span><b>{row["ratio"] * 100:.2f}%</b></div>'
        for row in sorted(summary, key=lambda item: item["value"], reverse=True)
    )
    center = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
    overlay_bounds = raster_bounds or bounds
    leaflet_overlay_bounds = [[overlay_bounds[1], overlay_bounds[0]], [overlay_bounds[3], overlay_bounds[2]]]
    class_overlay_name = class_overlay_png.name if class_overlay_png and class_overlay_png.exists() else None
    ndvi_overlay_name = ndvi_overlay_png.name if ndvi_overlay_png and ndvi_overlay_png.exists() else None

    crop_legend = f'<div class="legend-row"><i style="background:#20d8d2"></i><span>{crop_label or "作物"}</span><b>SHP</b></div>'

    html = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css" />
  <script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
  <style>
    html, body, #map {{ height: 100%; margin: 0; background: #101510; }}
    .leaflet-container {{ background: #101510; font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; }}
    .leaflet-control-layers, .leaflet-control-zoom {{
      border: 1px solid rgba(31,47,38,.16) !important;
      background: rgba(255,255,255,.94) !important;
      color: #17231c !important;
      box-shadow: 0 10px 28px rgba(25,38,30,.12) !important;
    }}
    .leaflet-control-layers label {{ color: #17231c; font: 12px system-ui, sans-serif; }}
    .leaflet-control-zoom a {{ background: #fff !important; color: #2f7d4d !important; border-color: rgba(31,47,38,.12) !important; }}
    .legend {{
      position: absolute;
      z-index: 500;
      right: 12px;
      bottom: 16px;
      width: 168px;
      border: 1px solid rgba(31,47,38,.14);
      border-radius: 8px;
      padding: 10px;
      background: rgba(255,255,255,.94);
      color: #17231c;
      font: 12px system-ui, sans-serif;
      backdrop-filter: blur(8px);
    }}
    .legend strong {{ display: block; margin-bottom: 8px; font-size: 13px; }}
    .legend-row {{ display: grid; grid-template-columns: 14px 1fr auto; gap: 8px; align-items: center; margin-top: 7px; }}
    .legend-row i {{ width: 12px; height: 12px; border-radius: 3px; }}
    .legend-row b {{ color: #64736a; font: 700 11px ui-monospace, SFMono-Regular, Menlo, monospace; }}
    .empty-note {{
      position: absolute;
      z-index: 500;
      left: 50%;
      top: 50%;
      transform: translate(-50%, -50%);
      border: 1px solid rgba(31,47,38,.14);
      border-radius: 8px;
      padding: 14px 16px;
      background: rgba(255,255,255,.94);
      color: #17231c;
      font: 13px system-ui, sans-serif;
      max-width: 320px;
      text-align: center;
    }}
  </style>
</head>
<body>
  <div id="map"></div>
  <div class="legend" id="legend"><strong>长势等级</strong>{legend_rows}</div>
  <script>
    const boundaryData = {json.dumps(boundary_data, ensure_ascii=False)};
    const growthLegendHtml = `<strong>长势等级</strong>{legend_rows}`;
    const cropLegendHtml = `<strong>作物图层</strong>{crop_legend}`;
    const osm = L.tileLayer('https://tile.openstreetmap.org/{{z}}/{{x}}/{{y}}.png', {{
      maxZoom: 20,
      maxNativeZoom: 19,
      attribution: '&copy; OpenStreetMap contributors'
    }});
    const satellite = L.tileLayer(
      'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{{z}}/{{y}}/{{x}}',
      {{
        maxZoom: 20,
        maxNativeZoom: 17,
        attribution: 'Esri World Imagery'
      }}
    );
    const map = L.map('map', {{
      zoomControl: true,
      preferCanvas: true,
      layers: [satellite]
    }}).setView({json.dumps(center)}, 13);
    const rasterBounds = {json.dumps(leaflet_overlay_bounds)};
    const overlays = {{}};
    const classOverlayName = {json.dumps(class_overlay_name, ensure_ascii=False)};
    const ndviOverlayName = {json.dumps(ndvi_overlay_name, ensure_ascii=False)};
    let ndviClassLayer = null;
    let ndviRawLayer = null;
    if (classOverlayName) {{
      ndviClassLayer = L.imageOverlay(classOverlayName, rasterBounds, {{
        opacity: 0.72,
        alt: 'NDVI classified raster from current calculation'
      }}).addTo(map);
      overlays['NDVI 长势覆盖'] = ndviClassLayer;
    }}
    if (ndviOverlayName) {{
      ndviRawLayer = L.imageOverlay(ndviOverlayName, rasterBounds, {{
        opacity: 0.62,
        alt: 'NDVI raster from current calculation'
      }});
      overlays['NDVI 原值覆盖'] = ndviRawLayer;
    }}
    const cropLayer = L.geoJSON(boundaryData, {{
      style: {{ color: '#00d5d1', weight: 1.5, fillColor: '#20d8d2', fillOpacity: 0.66 }}
    }});
    const boundary = L.geoJSON(boundaryData, {{
      style: {{ color: '#00f5ff', weight: 2, fillColor: '#00f5ff', fillOpacity: 0 }}
    }}).addTo(map);
    overlays['作物图层 SHP'] = cropLayer;
    overlays['地块边界'] = boundary;
    if (!classOverlayName && !ndviOverlayName) {{
      const note = document.createElement('div');
      note.className = 'empty-note';
      note.textContent = '当前任务没有可展示的本地栅格图层，请查看预览图或下载 TIF。';
      document.body.appendChild(note);
    }}
    const baseLayers = {{
      '卫星底图': satellite,
      'OSM 底图': osm
    }};
    L.control.layers(baseLayers, overlays, {{ collapsed: false }}).addTo(map);
    map.on('overlayadd overlayremove', () => {{
      const legend = document.getElementById('legend');
      if (!legend) return;
      if (map.hasLayer(cropLayer) && (!ndviClassLayer || !map.hasLayer(ndviClassLayer))) {{
        legend.innerHTML = cropLegendHtml;
      }} else {{
        legend.innerHTML = growthLegendHtml;
      }}
    }});
    map.fitBounds(boundary.getBounds(), {{ padding: [28, 28], maxZoom: 17 }});
  </script>
</body>
</html>"""
    out_html.parent.mkdir(parents=True, exist_ok=True)
    out_html.write_text(html, encoding="utf-8")
    return out_html


def _format_report_date(value: str | None) -> str:
    if not value:
        return "-"
    try:
        year, month, day = value.split("-")
        return f"{int(year)}.{int(month)}.{int(day)}"
    except Exception:
        return value


def _replace_paragraph_text(paragraph: Any, text: str) -> None:
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def _insert_summary_table_after(
    paragraph: Any,
    rows: list[dict[str, Any]],
    class_breaks: list[float] | None = None,
) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = paragraph.part.document
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["长势等级", "面积(亩)", "占比", "像元数", "NDVI上界"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(31, 47, 38)

    for row in sorted(rows, key=lambda item: item["value"], reverse=True):
        cells = table.add_row().cells
        upper = "-"
        if class_breaks and row["value"] - 1 < len(class_breaks):
            upper = f"{float(class_breaks[row['value'] - 1]):.4f}"
        values = [
            str(row["label"]),
            f"{float(row['area_mu']):.2f}",
            f"{float(row['ratio']) * 100:.2f}%",
            f"{int(row['count']):,}",
            upper,
        ]
        for idx, value in enumerate(values):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    paragraph._p.addnext(table._tbl)


def _replace_docx_media(src_docx: Path, out_docx: Path, replacements: dict[str, Path]) -> None:
    tmp_docx = out_docx.with_suffix(".media.docx")
    with zipfile.ZipFile(src_docx, "r") as zin, zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            replacement = replacements.get(item.filename)
            if replacement and replacement.exists():
                data = replacement.read_bytes()
            zout.writestr(item, data)
    if out_docx.exists():
        out_docx.unlink()
    tmp_docx.replace(out_docx)


def create_docx_report(
    rows: list[dict[str, Any]],
    out_docx: Path,
    insurer: str,
    crop_label: str,
    total_area_mu: float,
    method: str,
    *,
    class_breaks: list[float] | None = None,
    start_date: str | None = None,
    end_date: str | None = None,
    ndvi_meta: dict[str, Any] | None = None,
    growth_map_png: Path | None = None,
    crop_map_png: Path | None = None,
) -> Path | None:
    try:
        from docx import Document
    except Exception:
        return None

    template = Path(__file__).resolve().parent.parent / "作物长势监测报告（模板）.docx"
    if not template.exists():
        doc = Document()
        doc.add_heading("作物长势监测报告", level=1)
        doc.add_paragraph(f"投保人：{insurer or '-'}")
        doc.add_paragraph(f"作物：{crop_label or '-'}")
        doc.add_paragraph(f"监测面积：{total_area_mu:.2f} 亩")
        doc.add_paragraph(f"NDVI 分级方法：{method}")
        _insert_summary_table_after(doc.paragraphs[-1], rows, class_breaks)
        out_docx.parent.mkdir(parents=True, exist_ok=True)
        doc.save(out_docx)
        return out_docx

    ndvi_meta = ndvi_meta or {}
    start_label = _format_report_date(start_date or ndvi_meta.get("start_date"))
    end_label = _format_report_date(end_date or ndvi_meta.get("end_date"))
    place = insurer or "监测"

    doc = Document(template)
    summary_inserted = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("（2025."):
            _replace_paragraph_text(para, f"（{start_label}-{end_label}）")
        elif text.startswith("投保人："):
            _replace_paragraph_text(para, f"投保人： {insurer or '-'}")
        elif "xxx地块的实际投保面积" in text:
            _replace_paragraph_text(para, f"（1）{place}地块的实际投保面积")
        elif text == "作物：n亩":
            _replace_paragraph_text(para, f"{crop_label or '作物'}：{float(total_area_mu):.2f}亩")
        elif "根据卫星遥感显示" in text:
            source = ndvi_meta.get("source_label") or "卫星遥感"
            image_count = ndvi_meta.get("image_count")
            image_text = f"，有效影像 {image_count} 景" if image_count is not None else ""
            _replace_paragraph_text(
                para,
                f"根据{source}显示，{place}地块的作物长势情况如下（分级方法：{method}{image_text}）：",
            )
            if not summary_inserted:
                _insert_summary_table_after(para, rows, class_breaks)
                summary_inserted = True
        elif text.startswith("遥感显示的种植作物长势图如下"):
            _replace_paragraph_text(
                para,
                "遥感显示的种植作物长势图如下：（深绿色区域为“优”、浅绿色区域为“良”，黄色区域为“中”，橙色区域为“一般”，红色区域为“差”）",
            )

    tmp_docx = out_docx.with_suffix(".text.docx")
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(tmp_docx)
    media_replacements: dict[str, Path] = {}
    if growth_map_png and growth_map_png.exists():
        media_replacements["word/media/image1.png"] = growth_map_png
    if crop_map_png and crop_map_png.exists():
        media_replacements["word/media/image2.png"] = crop_map_png
    if media_replacements:
        _replace_docx_media(tmp_docx, out_docx, media_replacements)
        tmp_docx.unlink(missing_ok=True)
    else:
        tmp_docx.replace(out_docx)
    return out_docx


def run_growth_analysis(
    *,
    task_id: str,
    src_tif: Path,
    boundary_path: Path,
    output_dir: Path,
    total_area_mu: float,
    crop_label: str,
    insurer: str,
    method: str = "jenks",
    n_classes: int = 5,
    boundary_crs: str | None = None,
    report_context: dict[str, Any] | None = None,
    url_prefix: str = "",
) -> dict[str, Any]:
    """Run the full NDVI analysis and return frontend-ready metadata."""
    output_dir.mkdir(parents=True, exist_ok=True)

    ndvi_clip = output_dir / "ndvi_clip.tif"
    class_tif = output_dir / "ndvi_class.tif"
    ndvi_preview = output_dir / "ndvi_preview.png"
    class_preview = output_dir / "ndvi_class_preview.png"
    ndvi_overlay = output_dir / "ndvi_overlay.png"
    class_overlay = output_dir / "ndvi_class_overlay.png"
    report_growth_map = output_dir / "report_ndvi_growth_map.png"
    report_crop_map = output_dir / "report_crop_layer_map.png"
    summary_csv = output_dir / "summary.csv"
    summary_json = output_dir / "summary.json"
    class_geojson = output_dir / "ndvi_class.geojson"
    boundary_geojson = output_dir / "boundary.geojson"
    map_html = output_dir / "map.html"
    report_docx = output_dir / "growth_report.docx"

    clip_meta = clip_raster_by_boundary(src_tif, boundary_path, ndvi_clip, output_dir, boundary_crs)
    boundary_path_wgs84, boundary_bounds = write_boundary_geojson(
        boundary_path,
        output_dir,
        boundary_geojson,
        boundary_crs,
    )
    samples, stats = sample_raster_values(ndvi_clip)
    class_breaks = compute_class_breaks(samples, method, n_classes, stats)
    count_map = write_classified_raster(ndvi_clip, class_tif, class_breaks)
    rows = summarize_classes(count_map, total_area_mu, n_classes)

    render_ndvi_preview(ndvi_clip, ndvi_preview)
    render_class_preview(class_tif, class_preview, n_classes)
    ndvi_overlay_bounds = write_ndvi_overlay_png(ndvi_clip, ndvi_overlay)
    class_overlay_bounds = write_class_overlay_png(class_tif, class_overlay, n_classes)
    write_summary_csv(rows, summary_csv)
    summary_json.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")

    vector_path = write_class_geojson(class_tif, class_geojson)
    map_path = create_interactive_map_html(
        out_html=map_html,
        boundary_geojson=boundary_path_wgs84,
        class_overlay_png=class_overlay,
        ndvi_overlay_png=ndvi_overlay,
        summary=rows,
        bounds=boundary_bounds,
        raster_bounds=class_overlay_bounds or ndvi_overlay_bounds,
        crop_label=crop_label,
    )
    create_report_map_images(
        boundary_geojson=boundary_path_wgs84,
        class_overlay_png=class_overlay,
        ndvi_overlay_png=ndvi_overlay,
        overlay_bounds=class_overlay_bounds or ndvi_overlay_bounds,
        out_growth_png=report_growth_map,
        out_crop_png=report_crop_map,
        summary=rows,
        crop_label=crop_label,
    )
    report_context = report_context or {}
    report_path = create_docx_report(
        rows,
        report_docx,
        insurer,
        crop_label,
        total_area_mu,
        method,
        class_breaks=class_breaks,
        start_date=report_context.get("start_date"),
        end_date=report_context.get("end_date"),
        ndvi_meta=report_context.get("ndvi_meta"),
        growth_map_png=report_growth_map,
        crop_map_png=report_crop_map,
    )

    def url(path: Path | None) -> str | None:
        if path is None:
            return None
        return f"{url_prefix}/{path.name}" if url_prefix else str(path)

    outputs = {
        "ndvi_clip_tif": url(ndvi_clip),
        "classified_tif": url(class_tif),
        "ndvi_preview_png": url(ndvi_preview),
        "class_preview_png": url(class_preview),
        "ndvi_overlay_png": url(ndvi_overlay),
        "class_overlay_png": url(class_overlay),
        "report_growth_map_png": url(report_growth_map),
        "report_crop_map_png": url(report_crop_map),
        "summary_csv": url(summary_csv),
        "summary_json": url(summary_json),
        "map_html": url(map_path),
        "boundary_geojson": url(boundary_path_wgs84),
        "class_geojson": url(vector_path),
        "report_docx": url(report_path),
    }

    return {
        "status": "success",
        "task_id": task_id,
        "method": method,
        "n_classes": n_classes,
        "total_area_mu": round(float(total_area_mu), 2),
        "valid_pixel_count": int(sum(count for value, count in count_map.items() if value > 0)),
        "class_breaks": [round(float(x), 6) for x in class_breaks],
        "raster": clip_meta,
        "summary": rows,
        "outputs": outputs,
        "message": "长势分析完成",
    }


def run_growth_analysis_from_boundary(
    *,
    task_id: str,
    boundary_path: Path,
    output_dir: Path,
    total_area_mu: float | None,
    crop_label: str,
    insurer: str,
    method: str = "jenks",
    n_classes: int = 5,
    boundary_crs: str | None = None,
    ndvi_source: str = "auto",
    start_date: str = "2025-08-30",
    end_date: str = "2025-09-15",
    max_cloud_pct: float = 30.0,
    gee_project_id: str | None = None,
    url_prefix: str = "",
) -> dict[str, Any]:
    """Run the one-file workflow: boundary upload -> auto NDVI -> analysis."""
    output_dir.mkdir(parents=True, exist_ok=True)
    area_mu = (
        round(float(total_area_mu), 2)
        if total_area_mu is not None and float(total_area_mu) > 0
        else calculate_boundary_area_mu(boundary_path, output_dir, boundary_crs)
    )
    auto_ndvi = output_dir / "auto_ndvi.tif"
    ndvi_meta = prepare_auto_ndvi_raster(
        boundary_path=boundary_path,
        work_dir=output_dir,
        out_tif=auto_ndvi,
        source_crs=boundary_crs,
        ndvi_source=ndvi_source,
        start_date=start_date,
        end_date=end_date,
        max_cloud_pct=max_cloud_pct,
        gee_project_id=gee_project_id,
    )
    result = run_growth_analysis(
        task_id=task_id,
        src_tif=auto_ndvi,
        boundary_path=boundary_path,
        output_dir=output_dir,
        total_area_mu=area_mu,
        crop_label=crop_label,
        insurer=insurer,
        method=method,
        n_classes=n_classes,
        boundary_crs=boundary_crs,
        report_context={
            "start_date": start_date,
            "end_date": end_date,
            "ndvi_meta": ndvi_meta,
        },
        url_prefix=url_prefix,
    )
    result["outputs"]["auto_ndvi_tif"] = f"{url_prefix}/auto_ndvi.tif" if url_prefix else str(auto_ndvi)
    result["raster"]["ndvi_source"] = ndvi_meta.get("source")
    result["raster"]["ndvi_source_label"] = ndvi_meta.get("source_label")
    result["raster"]["ndvi_meta"] = ndvi_meta
    if ndvi_meta.get("source") == "synthetic":
        result["message"] = "边界自动解译完成（使用本地模拟 NDVI）"
    else:
        result["message"] = "边界自动解译完成（使用 GEE Sentinel-2 NDVI）"
    return result
